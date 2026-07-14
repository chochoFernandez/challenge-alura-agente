"""Genera la base documental ficticia de NovaPay (fintech imaginaria).

Se versiona el generador y no solo los archivos, para que la base de conocimiento sea
reproducible: cualquiera puede regenerar los 6 documentos desde cero y obtener lo mismo.

Cubre los 6 formatos que pide el challenge, uno por dominio de negocio:

    politica_rrhh.pdf         PDF       RH
    politica_privacidad.docx  DOCX      Legal / Compliance
    tarifas_comisiones.csv    CSV       Financiero
    planes_productos.json     JSON      Comercial
    faq_transacciones.md      Markdown  Operacional
    ventas_2025.xlsx          XLSX      Datos

Uso:
    python scripts/generate_docs.py
"""

import csv
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.config import get_settings

EMPRESA = "NovaPay"


# --------------------------------------------------------------------------------------
# 1. PDF — Política de Recursos Humanos
# --------------------------------------------------------------------------------------
# Nota deliberada: la sección de descanso anual se llama "Licencia Anual Remunerada" y
# evita la palabra "vacaciones". Así la pregunta "¿cuántos días de vacaciones tengo?"
# solo se puede responder con búsqueda semántica, no con coincidencia de palabras.
# Es exactamente el caso que el enunciado del challenge pone como ejemplo.

POLITICA_RRHH = [
    ("titulo", "Política de Recursos Humanos"),
    ("subtitulo", f"{EMPRESA} S.A. — Versión 3.1 — Vigente desde el 01/01/2025"),
    ("h1", "1. Objetivo y alcance"),
    (
        "p",
        "El presente documento establece las condiciones laborales, licencias y beneficios "
        f"aplicables a todas las personas que integran {EMPRESA} S.A. bajo relación de "
        "dependencia, en cualquiera de sus sedes o bajo modalidad remota. Reemplaza y deja sin "
        "efecto la versión 2.4 de marzo de 2023.",
    ),
    ("h1", "2. Jornada laboral y trabajo remoto"),
    (
        "p",
        "La jornada laboral es de 40 horas semanales, distribuidas de lunes a viernes. El horario "
        "central de colaboración obligatoria es de 10:00 a 16:00; fuera de esa franja cada equipo "
        "administra su horario con autonomía.",
    ),
    (
        "p",
        f"{EMPRESA} adopta un esquema híbrido. Cada colaborador puede trabajar de forma remota "
        "hasta 3 días por semana, debiendo asistir presencialmente los 2 días restantes. Los "
        "equipos definen en conjunto qué días son presenciales, priorizando la coincidencia para "
        "las ceremonias de equipo. El trabajo 100% remoto se autoriza solo por excepción y "
        "requiere aprobación del líder directo y de Personas.",
    ),
    ("h1", "3. Licencia Anual Remunerada"),
    (
        "p",
        "Todo colaborador tiene derecho a 22 días hábiles de licencia anual remunerada por año "
        "calendario trabajado. Este derecho se devenga a razón de 1,83 días por mes efectivamente "
        "trabajado y se acredita en el portal de Personas el primer día de cada mes.",
    ),
    (
        "p",
        "La licencia puede fraccionarse en hasta 3 períodos por año, siempre que ninguno de ellos "
        "sea inferior a 5 días corridos. La solicitud debe cargarse en el portal con una "
        "anticipación mínima de 15 días corridos y requiere la aprobación del líder directo, quien "
        "dispone de 5 días hábiles para responder.",
    ),
    (
        "p",
        "Los días no utilizados pueden acumularse hasta 18 meses contados desde el cierre del "
        "período en que se devengaron; vencido ese plazo caducan sin derecho a compensación "
        "económica. Durante el primer año de trabajo, el colaborador puede tomar licencia una vez "
        "cumplidos 6 meses de antigüedad, de forma proporcional a lo devengado.",
    ),
    ("h1", "4. Licencias especiales"),
    (
        "p",
        "Además de la licencia anual, se otorgan las siguientes licencias especiales, todas ellas "
        "remuneradas y sin descuento de la licencia anual: matrimonio o unión convivencial, 10 días "
        "corridos; nacimiento o adopción, 15 días corridos para la persona no gestante y 120 días "
        "para la persona gestante; fallecimiento de familiar directo, 5 días corridos; fallecimiento "
        "de familiar de segundo grado, 2 días corridos; mudanza, 1 día por año; y examen académico, "
        "2 días por examen con un máximo de 10 días por año calendario.",
    ),
    ("h1", "5. Beneficios"),
    (
        "p",
        "La empresa ofrece cobertura de salud de primer nivel para el colaborador y su grupo "
        "familiar primario, seguro de vida obligatorio, y un bono anual por desempeño equivalente a "
        "entre 0,5 y 2 salarios mensuales según la evaluación de fin de año.",
    ),
    (
        "p",
        "Cada colaborador dispone de un presupuesto anual de capacitación de USD 800, aplicable a "
        "cursos, certificaciones, libros o conferencias vinculadas a su rol. El presupuesto no es "
        "acumulable entre años y se solicita mediante el formulario interno de Desarrollo "
        "Profesional, con aprobación del líder directo.",
    ),
    ("h1", "6. Onboarding"),
    (
        "p",
        "El proceso de incorporación abarca los primeros 30 días. Cada nueva persona recibe un "
        "buddy asignado de su equipo, un checklist de accesos y herramientas que debe completarse "
        "en la primera semana, y tres instancias de seguimiento con Personas: a los 7, 30 y 90 días.",
    ),
    ("h1", "7. Área responsable"),
    (
        "p",
        "La interpretación y actualización de esta política corresponde a la Dirección de Personas. "
        "Responsable: María Estévez, Directora de Personas. Consultas: rrhh@novapay.example. "
        "Las solicitudes se gestionan a través del portal interno de Personas.",
    ),
]


def generar_pdf_rrhh(destino: Path) -> None:
    doc = SimpleDocTemplate(
        str(destino),
        pagesize=A4,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        title="Política de Recursos Humanos",
        author=f"{EMPRESA} S.A.",
    )

    base = getSampleStyleSheet()
    estilos = {
        "titulo": ParagraphStyle("titulo", parent=base["Title"], fontSize=20, spaceAfter=6),
        "subtitulo": ParagraphStyle(
            "subtitulo", parent=base["Normal"], fontSize=10, textColor="#666666", spaceAfter=18
        ),
        "h1": ParagraphStyle("h1", parent=base["Heading1"], fontSize=13, spaceBefore=14, spaceAfter=6),
        "p": ParagraphStyle("p", parent=base["BodyText"], fontSize=10.5, leading=15, alignment=TA_JUSTIFY),
    }

    story = []
    for tipo, texto in POLITICA_RRHH:
        story.append(Paragraph(texto, estilos[tipo]))
        if tipo == "p":
            story.append(Spacer(1, 4))

    # Encabezado y pie repetidos en cada página: son ruido real que app/ingest.py
    # tiene que limpiar, igual que en un PDF corporativo de verdad.
    def decorar(canvas, documento):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor("#888888")
        canvas.drawString(2.5 * cm, A4[1] - 1.5 * cm, f"{EMPRESA} S.A. — Documento interno confidencial")
        canvas.drawString(2.5 * cm, 1.5 * cm, "Política de Recursos Humanos v3.1")
        canvas.drawRightString(A4[0] - 2.5 * cm, 1.5 * cm, f"Página {documento.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=decorar, onLaterPages=decorar)


# --------------------------------------------------------------------------------------
# 2. DOCX — Política de Privacidad y Protección de Datos
# --------------------------------------------------------------------------------------

POLITICA_PRIVACIDAD = [
    ("h0", "Política de Privacidad y Protección de Datos"),
    ("sub", f"{EMPRESA} S.A. — Versión 2.0 — Vigente desde el 15/03/2025"),
    ("h1", "1. Datos que recolectamos"),
    (
        "p",
        "Recolectamos datos de identificación (nombre, documento, fecha de nacimiento), datos de "
        "contacto (correo electrónico, teléfono, domicilio), datos financieros y transaccionales "
        "(movimientos, saldos, medios de pago asociados), datos biométricos utilizados "
        "exclusivamente para la validación de identidad al momento del alta, y datos técnicos de "
        "navegación (dirección IP, identificador de dispositivo, registros de acceso).",
    ),
    ("h1", "2. Finalidad y base legal"),
    (
        "p",
        "Los datos se tratan para la prestación del servicio financiero contratado, el cumplimiento "
        "de obligaciones legales y regulatorias en materia de prevención de lavado de activos, la "
        "prevención del fraude, y —solo mediando consentimiento expreso y revocable— el envío de "
        "comunicaciones comerciales.",
    ),
    ("h1", "3. Plazos de conservación"),
    (
        "p",
        "Los datos de transacciones y los registros de operaciones se conservan durante 10 años "
        "desde la finalización de la relación comercial, plazo impuesto por la normativa de "
        "prevención de lavado de activos y financiación del terrorismo. Este plazo no puede "
        "reducirse a pedido del titular.",
    ),
    (
        "p",
        "Los datos utilizados con fines de marketing se eliminan a los 24 meses de la baja de la "
        "cuenta o de la revocación del consentimiento, lo que ocurra primero. Los registros "
        "técnicos de navegación se conservan por 12 meses.",
    ),
    ("h1", "4. Derechos del titular de los datos"),
    (
        "p",
        "El titular puede solicitar en cualquier momento el acceso a sus datos personales, su "
        "rectificación cuando sean inexactos, su supresión cuando no exista obligación legal de "
        "conservarlos, la limitación u oposición al tratamiento, y la portabilidad de sus datos en "
        "un formato estructurado y de uso común.",
    ),
    (
        "p",
        "Las solicitudes se presentan escribiendo a privacidad@novapay.example. La empresa dispone "
        "de un plazo de 15 días hábiles para responder, prorrogable por 15 días hábiles adicionales "
        "cuando la complejidad del pedido lo justifique, notificando la prórroga al titular.",
    ),
    ("h1", "5. Medidas de seguridad"),
    (
        "p",
        "La información se cifra en reposo con AES-256 y en tránsito con TLS 1.3. El acceso a datos "
        "personales por parte del personal está restringido según el principio de mínimo privilegio "
        "y queda registrado en logs de auditoría inalterables conservados por 24 meses.",
    ),
    ("h1", "6. Notificación de incidentes"),
    (
        "p",
        "Ante un incidente de seguridad que comprometa datos personales, la empresa notifica a la "
        "autoridad de control dentro de las 72 horas de tomar conocimiento, y a los titulares "
        "afectados sin dilación indebida cuando el incidente suponga un riesgo alto para sus "
        "derechos.",
    ),
    ("h1", "7. Transferencias internacionales"),
    (
        "p",
        "Algunos proveedores de infraestructura se encuentran fuera del país. Toda transferencia "
        "internacional se realiza al amparo de cláusulas contractuales tipo y únicamente hacia "
        "jurisdicciones con nivel de protección adecuado.",
    ),
    ("h1", "8. Área responsable"),
    (
        "p",
        "El responsable del tratamiento es la Gerencia Legal y Compliance. Oficial de Protección de "
        "Datos: Ing. Rubén Cortés. Contacto: privacidad@novapay.example. Consultas de compliance: "
        "legal@novapay.example.",
    ),
]


def generar_docx_privacidad(destino: Path) -> None:
    doc = Document()
    doc.core_properties.title = "Política de Privacidad y Protección de Datos"
    doc.core_properties.author = f"{EMPRESA} S.A."

    for tipo, texto in POLITICA_PRIVACIDAD:
        if tipo == "h0":
            doc.add_heading(texto, level=0)
        elif tipo == "sub":
            p = doc.add_paragraph(texto)
            p.runs[0].italic = True
        elif tipo == "h1":
            doc.add_heading(texto, level=1)
        else:
            doc.add_paragraph(texto)

    doc.save(destino)


# --------------------------------------------------------------------------------------
# 3. CSV — Tarifas y comisiones
# --------------------------------------------------------------------------------------

TARIFAS = [
    ("Transferencia interna entre cuentas NovaPay", "Transferencias", "0", "ARS", "Sin límite de operaciones", "2025-01-01"),
    ("Transferencia a otro banco por CBU o CVU", "Transferencias", "0", "ARS", "Primeras 5 por mes bonificadas; a partir de la sexta, ARS 150 por operación", "2025-01-01"),
    ("Extracción en cajero automático de la red", "Efectivo", "850", "ARS", "Por operación. Plan Plus: 3 extracciones bonificadas por mes. Plan Business: 8 bonificadas", "2025-01-01"),
    ("Extracción en cajero fuera de la red", "Efectivo", "1200", "ARS", "Por operación, sin bonificaciones", "2025-01-01"),
    ("Mantenimiento mensual de cuenta", "Cuenta", "0", "ARS", "Bonificado en todos los planes", "2025-01-01"),
    ("Reposición de tarjeta física por pérdida o robo", "Tarjetas", "4500", "ARS", "Primera reposición del año bonificada en Plan Business y Enterprise", "2025-01-01"),
    ("Consumo con tarjeta en el exterior", "Tarjetas", "3.5", "% sobre el monto", "Se aplica sobre el monto convertido a pesos", "2025-01-01"),
    ("Adelanto de efectivo con tarjeta de crédito", "Tarjetas", "5.0", "% sobre el monto", "Mínimo ARS 2.000 por operación", "2025-01-01"),
    ("Cobro con código QR (comercios)", "Comercios", "0.8", "% sobre la venta", "Más IVA. Acreditación en 48 horas hábiles", "2025-01-01"),
    ("Liquidación inmediata de ventas (comercios)", "Comercios", "1.8", "% sobre la venta", "Más IVA. Acreditación en el acto", "2025-01-01"),
    ("Gestión de contracargo o disputa", "Disputas", "0", "ARS", "Sin costo para el usuario", "2025-01-01"),
    ("Rechazo de débito automático por saldo insuficiente", "Cuenta", "600", "ARS", "Por cada rechazo", "2025-01-01"),
    ("Emisión de certificado de saldos y movimientos", "Cuenta", "0", "ARS", "Descarga inmediata desde la app", "2025-01-01"),
]


def generar_csv_tarifas(destino: Path) -> None:
    with destino.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["concepto", "categoria", "monto", "unidad", "condiciones", "vigencia_desde"])
        w.writerows(TARIFAS)


# --------------------------------------------------------------------------------------
# 4. JSON — Planes y productos
# --------------------------------------------------------------------------------------

PLANES = {
    "empresa": EMPRESA,
    "documento": "Catálogo de planes y productos",
    "vigencia_desde": "2025-01-01",
    "moneda": "ARS",
    "area_responsable": {
        "nombre": "Gerencia Comercial",
        "contacto": "comercial@novapay.example",
        "responsable": "Lucía Ferraro",
    },
    "planes": [
        {
            "id": "basica",
            "nombre": "Cuenta Básica",
            "precio_mensual": 0,
            "publico": "Personas físicas",
            "limite_transferencia_diario": 500000,
            "limite_transferencia_mensual": 3000000,
            "extracciones_gratis_por_mes": 0,
            "cashback": "0%",
            "usuarios_incluidos": 1,
            "soporte": "Chat en la app, respuesta en hasta 48 horas hábiles",
            "beneficios": ["Transferencias internas ilimitadas sin costo", "Tarjeta virtual sin cargo"],
        },
        {
            "id": "plus",
            "nombre": "Cuenta Plus",
            "precio_mensual": 4900,
            "publico": "Personas físicas",
            "limite_transferencia_diario": 2000000,
            "limite_transferencia_mensual": 15000000,
            "extracciones_gratis_por_mes": 3,
            "cashback": "1% en gastronomía y transporte, tope ARS 15.000 por mes",
            "usuarios_incluidos": 1,
            "soporte": "Chat en la app, respuesta en hasta 12 horas hábiles",
            "beneficios": ["Tarjeta física sin cargo", "3 extracciones bonificadas por mes", "Descuentos en comercios adheridos"],
        },
        {
            "id": "business",
            "nombre": "Cuenta Business",
            "precio_mensual": 12900,
            "publico": "PyMEs y monotributistas",
            "limite_transferencia_diario": 10000000,
            "limite_transferencia_mensual": 100000000,
            "extracciones_gratis_por_mes": 8,
            "cashback": "0,5% sobre consumos con tarjeta corporativa",
            "usuarios_incluidos": 5,
            "soporte": "Ejecutivo de cuenta asignado, respuesta en hasta 4 horas hábiles",
            "beneficios": ["Cobros con QR", "Hasta 5 usuarios con permisos diferenciados", "Integración con facturación electrónica", "Reportes contables exportables"],
        },
        {
            "id": "enterprise",
            "nombre": "Cuenta Enterprise",
            "precio_mensual": None,
            "precio_nota": "A medida, según volumen. Requiere cotización comercial.",
            "publico": "Empresas con más de 50 empleados",
            "limite_transferencia_diario": None,
            "limite_transferencia_nota": "Sin límite preestablecido; se define por contrato.",
            "extracciones_gratis_por_mes": 20,
            "cashback": "Negociable por contrato",
            "usuarios_incluidos": 50,
            "soporte": "Ejecutivo dedicado y soporte telefónico 24/7",
            "beneficios": ["API de pagos", "Conciliación automática", "SLA de disponibilidad 99,9%", "Onboarding asistido"],
        },
    ],
}


def generar_json_planes(destino: Path) -> None:
    destino.write_text(json.dumps(PLANES, ensure_ascii=False, indent=2), encoding="utf-8")


# --------------------------------------------------------------------------------------
# 5. Markdown — FAQ de transacciones
# --------------------------------------------------------------------------------------

FAQ = f"""# Preguntas frecuentes sobre transacciones

**{EMPRESA} S.A. — Área Operacional — Actualizado: 01/06/2025**

## ¿Cuánto tarda en acreditarse una transferencia?

Las transferencias entre cuentas {EMPRESA} son inmediatas, las 24 horas, todos los días del año.

Las transferencias a otros bancos por CBU o CVU también son inmediatas en la mayoría de los casos.
Sin embargo, cuando el monto supera los ARS 2.000.000, la operación pasa por una revisión
automática de prevención de fraude que puede demorar hasta 24 horas hábiles. Si la revisión se
extiende, el usuario recibe una notificación en la app.

## ¿Cuál es el límite diario de transferencias?

El límite depende del plan contratado. Se puede consultar el detalle en el catálogo de planes.
El límite se renueva a las 00:00 hora local.

Para ampliar el límite de forma permanente hay que solicitarlo desde la app, en Configuración >
Límites, y completar la validación de identidad reforzada. La ampliación se resuelve en hasta 72
horas hábiles.

## Transferí dinero a un CBU equivocado, ¿qué hago?

Una transferencia acreditada no puede revertirse de forma unilateral: el dinero ya está en la cuenta
del destinatario. Lo que se puede hacer es iniciar una **gestión de devolución** desde la app, en
Movimientos > Detalle de la operación > Solicitar devolución, dentro de los 30 días corridos.

{EMPRESA} contacta a la entidad receptora, que a su vez contacta al titular de la cuenta destino.
La devolución depende de la conformidad de esa persona: no está garantizada. El plazo estimado de
gestión es de 10 días hábiles.

## ¿Cómo desconozco un consumo que no reconozco?

El desconocimiento de un consumo debe iniciarse dentro de los 30 días corridos contados desde la
fecha de cierre del resumen en el que aparece la operación. Pasado ese plazo, el reclamo no puede
gestionarse por la vía de contracargo.

El trámite se inicia en la app, en Movimientos > Detalle > Desconocer consumo. La gestión de
contracargo no tiene costo. Mientras dura la investigación, el monto se acredita de forma
provisoria; si el contracargo se resuelve en contra del usuario, el monto se vuelve a debitar.

## ¿Puedo programar una transferencia para una fecha futura?

Sí. Las transferencias programadas se pueden agendar con hasta 90 días de anticipación y se ejecutan
a las 09:00 de la fecha elegida. Si en ese momento no hay saldo suficiente, el sistema reintenta a
las 15:00 y a las 21:00 del mismo día; luego la operación se cancela y se notifica al usuario.

## ¿Qué pasa si mi cuenta queda con saldo negativo?

Las cuentas {EMPRESA} no permiten saldo negativo por transferencias. En el caso de los débitos
automáticos, si no hay saldo suficiente el débito se rechaza y se aplica el cargo por rechazo
correspondiente, detallado en el tarifario vigente.

## Horarios de atención

El soporte por chat en la app funciona todos los días de 08:00 a 22:00. Fuera de ese horario se
puede dejar el mensaje y se responde al inicio del día siguiente. Los planes Business y Enterprise
cuentan con canales de atención diferenciados.

## Área responsable

Operaciones y Soporte. Contacto: soporte@novapay.example. Responsable: Diego Alonso, Jefe de
Operaciones.
"""


def generar_md_faq(destino: Path) -> None:
    destino.write_text(FAQ, encoding="utf-8")


# --------------------------------------------------------------------------------------
# 6. XLSX — Ventas 2025
# --------------------------------------------------------------------------------------
# El RAG recupera fragmentos de texto: no sabe sumar 5.000 filas. Por eso la planilla trae
# el resumen mensual YA AGREGADO, con una fila por mes y plan. Así la pregunta "¿qué plan
# se vendió más en diciembre?" se responde con un fragmento que literalmente contiene el
# dato, en vez de exigirle al modelo una agregación que no puede hacer.
# Esta limitación queda declarada en el README.

MESES = [
    "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre",
]

# (mes, plan, altas_netas, ingresos_ars)
ALTAS = {
    "Cuenta Básica": [1820, 1650, 1910, 2050, 1780, 1690, 1740, 1880, 2100, 2240, 2380, 2510],
    "Cuenta Plus": [940, 1010, 1180, 1240, 1320, 1290, 1410, 1560, 1720, 1980, 2340, 3120],
    "Cuenta Business": [210, 240, 265, 280, 310, 295, 330, 355, 390, 430, 480, 560],
    "Cuenta Enterprise": [4, 3, 6, 5, 7, 6, 8, 9, 11, 12, 14, 18],
}
PRECIOS = {"Cuenta Básica": 0, "Cuenta Plus": 4900, "Cuenta Business": 12900, "Cuenta Enterprise": 145000}


def generar_xlsx_ventas(destino: Path) -> None:
    wb = Workbook()

    hoja = wb.active
    hoja.title = "Resumen Mensual"
    hoja.append(["mes", "plan", "altas_netas", "ingresos_recurrentes_ars"])
    for celda in hoja[1]:
        celda.font = Font(bold=True)

    for i, mes in enumerate(MESES):
        for plan, valores in ALTAS.items():
            altas = valores[i]
            hoja.append([mes, plan, altas, altas * PRECIOS[plan]])

    metricas = wb.create_sheet("Metricas Clave")
    metricas.append(["metrica", "valor", "unidad", "periodo"])
    for celda in metricas[1]:
        celda.font = Font(bold=True)

    total_altas = sum(sum(v) for v in ALTAS.values())
    filas_metricas = [
        ["Altas netas totales del año", total_altas, "cuentas", "2025"],
        ["Plan con más altas en el año", "Cuenta Básica", "plan", "2025"],
        ["Plan con más altas en diciembre", "Cuenta Plus", "plan", "Diciembre 2025"],
        ["Plan con mayor ingreso recurrente en diciembre", "Cuenta Plus", "plan", "Diciembre 2025"],
        ["Churn mensual promedio", 2.4, "%", "2025"],
        ["Ingreso promedio por usuario (ARPU)", 3180, "ARS", "2025"],
        ["Mes con más altas del año", "Diciembre", "mes", "2025"],
        ["Tasa de conversión de Básica a Plus", 8.7, "%", "2025"],
    ]
    for fila in filas_metricas:
        metricas.append(fila)

    wb.save(destino)


# --------------------------------------------------------------------------------------

GENERADORES = [
    ("politica_rrhh.pdf", generar_pdf_rrhh),
    ("politica_privacidad.docx", generar_docx_privacidad),
    ("tarifas_comisiones.csv", generar_csv_tarifas),
    ("planes_productos.json", generar_json_planes),
    ("faq_transacciones.md", generar_md_faq),
    ("ventas_2025.xlsx", generar_xlsx_ventas),
]


def main() -> None:
    settings = get_settings()
    settings.docs_dir.mkdir(parents=True, exist_ok=True)

    print(f"Generando la base documental de {EMPRESA} en {settings.docs_dir}\n")
    for nombre, generador in GENERADORES:
        destino = settings.docs_dir / nombre
        generador(destino)
        print(f"  ✓ {nombre:28} {destino.stat().st_size:>7,} bytes")

    print(f"\n{len(GENERADORES)} documentos generados.")


if __name__ == "__main__":
    main()
